from docx import Document
import os,re,datetime
from openpyxl import load_workbook
import pandas as pd


items = {"s/f": {"Habilitación":"f.2", "Extension de Vida Util":"f.4","Control Periodico":"f.10"},
         "c/f": {"Habilitación":"f.3", "Extension de Vida Util":"f.5","Control Periodico":"f.11"}}

info = dict()


def read(path, cols, names,header, sheetname) -> pd.DataFrame:
    return pd.read_excel(path,
            usecols=cols, 
            names=names, 
            header=header,
            sheet_name=sheetname)


def numero_a_mes(mes) -> str:
    m = {
        'enero': "01",
        'febrero': "02",
        'marzo': "03",
        'abril': "04",
        'mayo': "05",
        'junio': "06",
        'julio': "07",
        'agosto': "08",
        'septiembre': "09",
        'octubre': "10",
        'noviembre': "11",
        'diciembre': "12"
        }
    aux = dict(zip(m.values(),m.keys()))
    return(aux[mes])


def crear_nota_presentacion(path,fecha):
    plantilla = Document(os.path.join(path, r"Templates\cartaSPA conjunta.docx"))
    ultima_fila = 15
    mes = numero_a_mes(datetime.datetime.strptime(fecha,"%d/%m/%Y").strftime("%m")).title()
    plantilla.paragraphs[0].runs[2].text = plantilla.paragraphs[0].runs[2].text.replace("{MES}", mes)
    for k in info.keys():
        parrafo = plantilla.paragraphs[ultima_fila]
        for run in parrafo.runs:
            texto = run.text
            texto = texto.replace('{MES}', info[k]["Denominacion"])
            texto = texto.replace('{DENOMINACION}', info[k]["Denominacion"])
            texto = texto.replace('{FLUIDO}', info[k]["Fluido"])
            texto = texto.replace('{INTERNO}', k)
            texto = texto.replace('{PRESION}', str(info[k]["Presion"]))
            texto = texto.replace('{VOLUMEN}', str(info[k]["Volumen"]).replace("m3","").strip())
            run.text = texto
        ultima_fila += 1
    plantilla.save(os.path.join(path,r'Documentacion Generada\Nota de presentacion.docx'))
    crear_crono(path,fecha)


def crear_crono(path,fecha):
    lista = []
    for k in info.keys():
        lista.append([k, fecha,fecha,fecha,info[k]["Presion"].replace(".",","),info[k]["Volumen"].replace(".",",")])
    cols = ("Equipo","Inspeccion Visual","Control Dimensional","END","Presion","Volumen")
    df = pd.DataFrame(lista)
    df.columns = cols
    df.to_excel(os.path.join(path,f"Documentacion Generada\Crono.xlsx"), index=False)


def crear_caratula(path, denominacion, cliente, interno, presion,volumen,fluido,tarea,cf,filename,fecha,index,cantidad_equipos):
    if len(denominacion) > 25 and len(interno) > 25 and len(fluido)>25:
        aux = r"Templates\Caratula  Colegio Template - XL.docx"
    elif len(denominacion) > 25 and len(interno) > 25 or len(denominacion)>25 and len(fluido)>25 or len(interno)>25 and len(fluido)>25:
        aux = r"Templates\Caratula  Colegio Template - L.docx"
    elif len(denominacion) > 25 or len(interno) > 25 or len(fluido) > 25:
        aux = r"Templates\Caratula  Colegio Template - M.docx"
    else:
        aux = r"Templates\Caratula  Colegio Template.docx"
    plantilla = Document(os.path.join(path, aux))
    item = items["s/f"][tarea] if not cf else items["c/f"][tarea]
    mes = numero_a_mes(fecha.split("/")[1]).capitalize()
    anio = fecha.split("/")[2]
    for c,parrafo in enumerate(plantilla.paragraphs):
        if c==0 or c>12:
            for run in parrafo.runs:
                texto = run.text
                texto = texto.replace('{MES}', mes)
                texto = texto.replace('{AÑO}', f" {anio}")
                texto = texto.replace('{DENOMINACION}', str(denominacion).replace("T-",""))
                texto = texto.replace('{FLUIDO}', fluido)
                texto = texto.replace('{CLIENTE}', cliente.upper())
                texto = texto.replace('{TAREA}', tarea.title())
                texto = texto.replace('{INTERNO}', interno.replace("T-",""))
                texto = texto.replace('{PRESION}', str(presion))
                texto = texto.replace('{VOLUMEN}', str(volumen))
                texto = texto.replace('{ITEM}', item)
                run.text = texto
    info[interno] = {"Denominacion":"","Presion":0.0,"Volumen":0,"Fluido":""}
    info[interno]["Fluido"] = fluido
    info[interno]["Presion"] = presion
    info[interno]["Volumen"] = volumen.replace("m3","").strip()
    info[interno]["Denominacion"] = denominacion
    if index == cantidad_equipos-1:
        crear_nota_presentacion(path, fecha)
    plantilla.save(os.path.join(path,f'Documentacion Generada\CARATULAS\{index}-Caratula de colegio {filename}.docx'))
    plantilla.save(os.path.join(path,f'Documentacion Generada\{index}-1-Caratula de colegio {filename}.docx'))


def crear_memoria_desc(path, fecha,
            comitente,
            fabricante,
            modelo,
            id,
            año_de_fabricacion,
            fluido,
            volumen,
            norma,
            material,
            presion,
            diametro_envolvente,
            longitud_envolvente,
            emm_envolvente,
            cabezal,
            emm_cabezal,
            aislacion,
            tarea,
            temperatura_trabajo,
            n_acta,
            prox_control,
            index,
            cantidad_equipos):
    if _memoria_xl(len(comitente), len(fabricante), len(id), len(fluido)):
        plantilla = Document(os.path.join(path, r"Templates\Memoria Descriptiva Template - L.docx"))
    elif len(comitente)>40 or len(id)>40:
        plantilla = Document(os.path.join(path, r"Templates\Memoria Descriptiva Template - M.docx"))
    else:
        plantilla = Document(os.path.join(path, r"Templates\Memoria Descriptiva Template.docx"))
    if aislacion == "x":
        aislacion = "POSEE"
    else:
        aislacion = "NO POSEE"
    formulario_f = read(os.path.join(path, "Formulario F.xlsx"), "A,G,H", ["Denominacion","N_acta","Vto"], 0, "page 1")
    den = formulario_f.loc[formulario_f["N_acta"]==n_acta, "Denominacion"].iloc[0]
    disposicion = modelo if modelo.lower() in {"vertical","horizontal"} else "-"
    for c,parrafo in enumerate(plantilla.paragraphs):
        if c>3:
            for run in parrafo.runs:
                texto = run.text
                if "{" not in texto:
                    continue
                texto = texto.replace('{COMITENTE}', comitente)
                texto = texto.replace('{ID_INTERNA}', id)
                texto = texto.replace('{FABRICANTE}', fabricante)
                texto = texto.replace('{AÑO_FABRICACION}', año_de_fabricacion)
                texto = texto.replace('{DISPOSICION}', disposicion)
                texto = texto.replace('{SUSTENTACION}', "-")
                texto = texto.replace('{FLUIDO}', fluido)
                texto = texto.replace('{PRESION}', str(presion))
                texto = texto.replace('{TEMPERATURA_TRABAJO}', temperatura_trabajo)
                texto = texto.replace('{VOLUMEN}', str(volumen))
                texto = texto.replace('{ENVOLVENTE}', str(diametro_envolvente))
                texto = texto.replace('{LONGITUD_ENVOLVENTE}', str(longitud_envolvente))
                texto = texto.replace('{CABEZAL}', cabezal)
                texto = texto.replace('{MATERIAL}', material)
                texto = texto.replace('{AISLACION}', aislacion)
                texto = texto.replace('{NORMA}', norma)
                texto = texto.replace('{PRESION_MAX}', str(presion))
                texto = texto.replace('{ESP_CABEZAL}', str(emm_cabezal) if float(emm_cabezal) != 0 else "-")
                texto = texto.replace('{ESP_ENV}', str(emm_envolvente))
                run.text = texto

    filename = id.replace("/","-")
    filename = [s for s in filename if s not in "\/:*?<>|\n"]
    filename = "".join(filename)
    prox_control = formulario_f.loc[formulario_f["N_acta"]==n_acta, "Vto"].iloc[0]
    crear_caratula(path, den, comitente, id, presion, volumen, fluido, tarea, False, filename,fecha,index,cantidad_equipos)
    if re.search("ASME",norma) or norma == "":
        crear_memoria_asme(path, comitente,id,fecha,emm_envolvente,emm_cabezal,diametro_envolvente, presion,den,fabricante,material,cabezal,filename,prox_control,index,disposicion)
    else:
        crear_memoria_merkblatter(path, comitente,id,fecha,emm_envolvente,emm_cabezal,diametro_envolvente, presion,den,fabricante,material,cabezal,temperatura_trabajo,filename,prox_control,index,disposicion)
    # plantilla.save(os.path.join(path,f'Documentacion Generada\MEMORIAS DESCRIPTIVAS\{index}-Memoria Descriptiva {filename}.docx'))
    plantilla.save(os.path.join(path,f'Documentacion Generada\{index}-2-Memoria Descriptiva {filename}.docx'))


def _memoria_xl(a,b,c,d):
    return a>40 and b>40 or a>40 and c>40 or a>40 and d>40 or b>40 and c>40 or b>40 and d>40 or c>40 and d>40


def crear_memoria_asme(path,cliente,equipo,fecha,emm_envolvente,emm_cabezal,diametro_env, presion, denominacion, fabricante,material,cabezal,filename,prox_control,index,disposicion_equipo=""):
    if cabezal.lower() == "semiesferico":
        wb = load_workbook(os.path.join(path,"Templates\Memoria de Calculo ASME Template SEMIESF.xlsx"))
    else:
        wb = load_workbook(os.path.join(path,"Templates\Memoria de Calculo ASME Template.xlsx"))
    df = read(os.path.join(path, "Templates\Clientes Ubicacion.xlsx"), "A:F", ["Cliente","Calle","Nº","Partido","Localidad", "CP"], 0, "Sheet1")
    
    tensiones = read(os.path.join(path, "varios\Tensiones.xlsx"), "B:D", ["Norma","Material","Tension"], 0, "Sheet1")
    aux = df.to_dict()
    tension = tensiones.loc[tensiones["Material"] == material.lower(), "Tension"].iloc[0] if not tensiones.loc[tensiones["Material"] == material.lower(), "Tension"].empty else 2450
    index_cliente = dict(zip(aux["Cliente"].values(),aux["Cliente"].keys()))[cliente] if cliente in set(aux["Cliente"].values()) else ""
    hoja = wb.active
    if index_cliente:
        hoja["G4"] = cliente 
        hoja["G5"] = aux["Calle"][index_cliente]
        hoja["G6"] = aux["Localidad"][index_cliente]
        hoja["G7"] = aux["Partido"][index_cliente]
        hoja["N5"] = aux["Nº"][index_cliente]
        hoja["N6"] = aux["CP"][index_cliente]
    else:
        hoja["G4"] = "Agregar informacion de cliente" 
    
    disposicion = disposicion_equipo if disposicion_equipo != "-" else "Vertical"
    hoja["AF5"] = equipo
    hoja["AF7"] = fecha
    emm_cabezal = float(emm_cabezal) if float(emm_cabezal) != 0 else float(emm_envolvente)
    hoja["D11"] = emm_cabezal
    hoja["F11"] = float(emm_envolvente)
    hoja["H11"] = emm_cabezal
    hoja["AE12"] =float(presion)
    hoja["AE13"] = float(diametro_env)
    hoja["AE15"] = float(tension)
    hoja["AE16"] = float(tension)
    hoja["M22"] = denominacion
    hoja["M23"] = fabricante
    hoja["M24"] = material
    hoja["M26"] = disposicion
    hoja["M30"] = cabezal
    hoja["AD39"] = prox_control
    wb.save(os.path.join(path,f"Documentacion Generada\{index}-3-Memoria de Calculo ASME-{filename}.xlsx"))
    # wb.save(os.path.join(path,f"Documentacion Generada\MEMORIAS DE CALCULO\{index}-Memoria de Calculo ASME-{filename}.xlsx"))


def crear_memoria_merkblatter(path,cliente,equipo,fecha,emm_envolvente,emm_cabezal,diametro_env, presion, denominacion, fabricante,material,cabezal,temperatura_trabajo,filename,prox_control,index,disposicion_equipo=""):
    
    if cabezal.lower() == "semiesferico":
        wb = load_workbook(os.path.join(path,"Templates\Memoria de Calculo MERKBLATTER Template SEMIESF.xlsx"))
    else:
        wb = load_workbook(os.path.join(path,"Templates\Memoria de Calculo MERKBLATTER Template.xlsx"))
    df = read(os.path.join(path, "Templates\Clientes Ubicacion.xlsx"), "A:F", ["Cliente","Calle","Nº","Partido","Localidad", "CP"], 0, "Sheet1")
    aux = df.to_dict()
    tensiones = read(os.path.join(path, "varios\Tensiones.xlsx"), "B:D", ["Norma","Material","Tension"], 0, "Sheet1")
    tension = tensiones.loc[tensiones["Material"] == material.lower(), "Tension"].iloc[0] if not tensiones.loc[tensiones["Material"] == material.lower(), "Tension"].empty else 1230.66
    index_cliente = dict(zip(aux["Cliente"].values(),aux["Cliente"].keys()))[cliente] if cliente in set(aux["Cliente"].values()) else ""
    hoja = wb.active
    if index_cliente:
        hoja["G4"] = cliente 
        hoja["G5"] = aux["Calle"][index_cliente]
        hoja["G6"] = aux["Localidad"][index_cliente]
        hoja["G7"] = aux["Partido"][index_cliente]
        hoja["N5"] = aux["Nº"][index_cliente]
        hoja["N6"] = aux["CP"][index_cliente]
    else:
        hoja["G4"] = "Agregar informacion de cliente" 
    
    disposicion = disposicion_equipo if disposicion_equipo != "-" else "Vertical"
    hoja["AF5"] = equipo
    hoja["AF7"] = fecha
    emm_cabezal = float(emm_cabezal) if float(emm_cabezal) != 0 else float(emm_envolvente)
    hoja["D11"] = emm_cabezal
    hoja["F11"] = float(emm_envolvente)
    hoja["H11"] = emm_cabezal
    hoja["AE11"] = float(diametro_env)
    hoja["AE13"] = round(float(presion)*0.9806,1)
    hoja["AE14"] = float(tension)
    hoja["AE19"] = temperatura_trabajo
    hoja["M22"] = denominacion
    hoja["M23"] = fabricante
    hoja["M24"] = material
    hoja["M26"] = disposicion
    hoja["M30"] = cabezal
    hoja["AD39"] = prox_control
    wb.save(os.path.join(path,f"Documentacion Generada\{index}-3-Memoria de Calculo MERKBLATTER-{filename}.xlsx"))
    # wb.save(os.path.join(path,f"Documentacion Generada\MEMORIAS DE CALCULO\{index}-Memoria de Calculo MERKBLATTER-{filename}.xlsx"))


