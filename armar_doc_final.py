from pypdf import PdfMerger, PdfReader, PdfWriter
from docx import Document
import comtypes.client
import os, re, time,datetime

def main(path):
    merger = PdfWriter()
    
    files, sorted_files = os.listdir(path), []
    aux_merger = PdfWriter()
    for f in files:
        if f.endswith(".pdf"):
            aux = f.split("-",maxsplit=1)
            sorted_files.append((int(aux[0]),aux[1]))
    sorted_files.sort(key= lambda x: x[0])
    files = ["".join([str(x[0]),"-",x[1]]) for x in sorted_files]

    merger.append(os.path.join(path,files[0]))
    merger.append(os.path.join(path,files[1]))
    merger.append(os.path.join(path,files[-2]))
    merger.append(os.path.join(path,files[3]))
    merger.append(os.path.join(path,files[4]))
    merger.append(os.path.join(path,files[2]))
    # merger.merge(position=0, fileobj=os.path.join(path,files[1]))
    reader = PdfReader(os.path.join(path,files[8]))
    pre_hora = datetime.datetime.now().time()
    pre = time.time()
    generar_caratulas_general(path,reader)
    merge_caratulas_documentacion(path,files[8],aux_merger)
    merge_equipos_actas(path, PdfReader(os.path.join(path,"Final","equipos.pdf")),PdfReader(os.path.join(path,files[5])))
    merger.append(os.path.join(path,"Final","equipos_actas.pdf"))
    merger.append(os.path.join(path,files[7]))
    merger.append(os.path.join(path,files[10]))
    
    output = open(os.path.join(path,"Final","Doc Final.pdf"),"wb")
    merger.write(output)
    elapsed = time.time() - pre
    post_hora = datetime.datetime.now().time()
    print(f"Inicio:{pre_hora}\nFin:{post_hora}\nTotal:{elapsed}")


def generar_caratulas_general(path,reader:PdfReader):
    for i in range(0,len(reader.pages),3):
        text = reader.pages[i].extract_text()
        tarea = text[re.search("Encuadre de la tarea:",text).span()[1]:re.search("Item",text).span()[0]].strip()
        if tarea == "Control Periodico":
            plantilla = Document(os.path.join(path, r"CaratulaTemplateC.docx"))
        elif tarea == "Habilitación":
            plantilla = Document(os.path.join(path, r"CaratulaTemplateH.docx"))
        else:
            raise Exception("Caratula template not found")
        den = text[re.search("Denominación del equipo:",text).span()[1]:re.search("Fluido",text).span()[0]].strip().replace("\n", " ")
        interno = text[re.search("Nº interno:",text).span()[1]:re.search("Encuadre",text).span()[0]].strip().replace("\n", " ")
        interno = interno.split("-",maxsplit=1)
        interno = f"{interno[0]}\n{interno[1]}"
        plantilla.paragraphs[1].runs[0].text = plantilla.paragraphs[1].runs[0].text.replace('{DENOMINACION}', den)
        # plantilla.tables[0].cell(0,0).text = plantilla.tables[0].cell(0,0).text.replace('{EQUIPO}', interno)
        plantilla.paragraphs[5].runs[0].text = plantilla.paragraphs[5].runs[0].text.replace('{EQUIPO}', interno)
        plantilla.save(os.path.join(path,"Dump",f'{i}.docx'))


def merge_caratulas_documentacion(path,path_documentacion,aux_merger):
    sorted_files = []
    for f in os.listdir(os.path.join(path,"Dump")):
        if f.endswith(".docx"):
            aux = f.split(".",maxsplit=1)
            sorted_files.append((int(aux[0]),aux[1]))
    sorted_files.sort(key= lambda x: x[0])
    sorted_files = ["".join([str(x[0]),".",x[1]]) for x in sorted_files]
    for f in sorted_files:
        if f.endswith(".docx"):
            wdFormatPDF = 17
            word = comtypes.client.CreateObject('Word.Application')
            doc = word.Documents.Open(os.path.join(path,"Dump",f))
            # doc = word.Documents.Open(os.path.join(path,"Dump",f'{i}.docx'))
            output_name = f.replace(".docx", "")
            index = int(output_name)
            doc.SaveAs(os.path.join(path,"Dump",f'{output_name}.pdf'), FileFormat=wdFormatPDF)
            doc.Close()
            word.Quit()
            
            aux_merger.append(os.path.join(path,"Dump",f'{output_name}.pdf'))
            aux_merger.append(fileobj=os.path.join(path,path_documentacion),pages=(index,index+3))

    output = open(os.path.join(path,"Final","equipos.pdf"),"wb")
    aux_merger.write(output)


def merge_equipos_actas(path,reader_equipos:PdfReader,reader_actas:PdfReader):
    merger_equipos_actas = PdfWriter()
    for i in range(0,len(reader_equipos.pages),4):
        text = reader_equipos.pages[i+1].extract_text()
        print(text)
        interno = text[re.search("Nº interno:",text).span()[1]:re.search("Encuadre",text).span()[0]].strip()
        for j in range(0,len(reader_actas.pages),2):
            texto_acta = reader_actas.pages[j].extract_text()
            interno_acta = texto_acta[re.search("Identificacion Interna:",texto_acta).span()[1]:re.search("Registro Habilitante",texto_acta).span()[0]].replace("\n", " ").strip()
            if interno == interno_acta:
                print(f"Index:{i}-Equipo:{interno}-Acta:{interno_acta}")
                merger_equipos_actas.append(reader_equipos,pages=(i,i+4))
                merger_equipos_actas.append(reader_actas,pages=(j,j+2))
                break
    output = open(os.path.join(path,"Final","equipos_actas.pdf"),"wb")
    merger_equipos_actas.write(output)


main(r"C:\Users\Cristian\Documents\GO\DocFinalTest")